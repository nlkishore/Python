from docx import Document
#sfrom docx.shared import Pt

def is_heading(paragraph):
    """Check if the paragraph is a heading."""
    return paragraph.style.name.startswith("Heading")

def copy_tables_with_headings(source_doc, target_doc):
    src_doc = Document(source_doc)
    tgt_doc = Document()

    paragraphs = src_doc.paragraphs
    tables = src_doc.tables
    table_index = 0

    for i, para in enumerate(paragraphs):
        if table_index < len(tables) and para._element.getnext() is not None:
            next_element = para._element.getnext()
            if next_element.tag.endswith("tbl"):  # Check if next element is a table
                # Copy heading before table
                if is_heading(para):
                    new_para = tgt_doc.add_paragraph(para.text)
                    new_para.style = para.style  # Preserve style
                
                # Copy the table
                table = tables[table_index]
                new_table = tgt_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                new_table.style = table.style  # Preserve table style

                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        new_table.cell(r_idx, c_idx).text = cell.text  # Copy text

                table_index += 1  # Move to the next table

    # Save the new document
    tgt_doc.save(target_doc)

# Example usage
source_file = "source.docx"
target_file = "output.docx"
copy_tables_with_headings(source_file, target_file)

print("Tables and headings copied successfully!")
