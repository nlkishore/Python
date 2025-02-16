from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def iter_block_items(parent):
    """
    Yield each paragraph and table child within parent, in document order.
    """
    for child in parent.element.body.iterchildren():
        if child.tag.endswith('}p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)

# Load the original Word document
original_doc = Document('your_document.docx')

# Create a new Word document
new_doc = Document()

# List to keep track of all block items from the original document
block_items = list(iter_block_items(original_doc))

# Specify how many lines (paragraphs) above the table you want to copy
num_lines = 3

# Iterate over the block items
for idx, item in enumerate(block_items):
    if isinstance(item, Table):
        # Collect paragraphs above the table
        lines_above = []
        for i in range(1, num_lines + 1):
            prev_idx = idx - i
            if prev_idx >= 0:
                prev_item = block_items[prev_idx]
                if isinstance(prev_item, Paragraph):
                    text = prev_item.text.strip()
                    if text:
                        lines_above.insert(0, prev_item)
            else:
                break  # Reached the start of the document

        # Copy the collected paragraphs to the new document
        for para in lines_above:
            new_paragraph = new_doc.add_paragraph()
            new_paragraph.style = para.style  # Copy paragraph style
            new_paragraph.alignment = para.alignment  # Copy alignment
            # Copy runs and formatting
            for run in para.runs:
                new_run = new_paragraph.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
                new_run.font.color.rgb = run.font.color.rgb
                new_run.font.highlight_color = run.font.highlight_color

        # Add a blank paragraph for spacing (optional)
        new_doc.add_paragraph()

        # Copy the table to the new document
        # Create a new table with the same number of rows and columns
        new_table = new_doc.add_table(rows=0, cols=0)
        new_table.style = item.style  # Copy table style

        # Copy each row from the original table
        for row in item.rows:
            new_row = new_table.add_row()
            # Ensure the new row has the correct number of cells
            while len(new_row.cells) < len(row.cells):
                new_row.add_cell()
            for idx_cell, cell in enumerate(row.cells):
                new_cell = new_row.cells[idx_cell]
                # Clear any default content
                new_cell.text = ''
                # Copy cell content and formatting
                for paragraph in cell.paragraphs:
                    new_paragraph = new_cell.add_paragraph()
                    new_paragraph.style = paragraph.style
                    new_paragraph.alignment = paragraph.alignment
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.size = run.font.size
                        new_run.font.name = run.font.name
                        new_run.font.color.rgb = run.font.color.rgb
                        new_run.font.highlight_color = run.font.highlight_color
                # Remove the initial empty paragraph added by default
                if new_cell.paragraphs[0].text == '':
                    p = new_cell.paragraphs[0]._element
                    p.getparent().remove(p)
                    p._p = p._element = None

        # Add a blank paragraph after the table (optional)
        new_doc.add_paragraph()

# Save the new document
new_doc.save('extracted_tables.docx')
