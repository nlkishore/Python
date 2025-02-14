# First, make sure you've installed python-docx:
# pip install python-docx

from docx import Document

# Load your Word document
doc = Document('your_document.docx')

def readSingleTable():
    # Access the first table in the document
    table = doc.tables[0]

    # Initialize a list to store the table data
    data = []

    # Iterate over the rows in the table
    for row in table.rows:
        # Extract text from each cell in the row
        row_data = [cell.text for cell in row.cells]
        data.append(row_data)
    # Now 'data' contains all the table content
    for row in data:
        print(row)

def readAllTables():
    doc = Document('your_document.docx')

    # Iterate over all tables in the document
    for idx, table in enumerate(doc.tables):
        print(f"Processing Table {idx + 1}:")
        data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            data.append(row_data)
        for row in data:
            print(row)
        print("\n")  # Add a newline for better readability between tables

def processTableDataIntodf(data):
    import pandas as pds
    df = pd.DataFrame(data[1:], columns=data[0])  # Assuming the first row is the header
    print(df)

def readSpecificTable():
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    for row in table.rows:
        for cell in row.cells:
            cell.text = 'Sample Text'

    doc.save('new_document.docx')

def createNewTable():
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    for row in table.rows:
        for cell in row.cells:
            cell.text = 'Sample Text'

    doc.save('new_document.docx')


